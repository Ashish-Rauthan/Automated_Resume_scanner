"""
utils.py
--------
Resume text extraction supporting both PDF (.pdf) and Word (.docx) formats.

Public API:
  extract_text_from_bytes(file_bytes, filename)
      -> auto-detects format from magic bytes (not just extension)
      -> returns (text, error) tuple, never raises

  validate_resume_file(filename, file_bytes, max_size_mb)
      -> accepts .pdf and .docx
      -> returns None (valid) or error string

  clean_text(raw)
      -> shared normalisation used by both extractors

validate_pdf_file is kept as an alias for backward compatibility
with main.py which still calls validate_pdf_file().
"""

import io
import re
import logging
from typing import Optional, Tuple

import pdfplumber
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_SIZE_MB = 5.0


# ── Text cleaning (shared) ────────────────────────────────────────────────────

def clean_text(raw: str) -> str:
    """
    Normalise extracted resume text regardless of source format.
    Handles ligatures, smart punctuation, whitespace, blank lines.
    """
    if not raw:
        return ""

    replacements = {
        "\ufb01": "fi",  "\ufb02": "fl",  "\ufb00": "ff",
        "\ufb03": "ffi", "\ufb04": "ffl",
        "\u2022": "-",   "\u2023": "-",   "\u25cf": "-",
        "\u2013": "-",   "\u2014": "-",
        "\u2019": "'",   "\u2018": "'",
        "\u201c": '"',   "\u201d": '"',
        "\u00a0": " ",   "\u2026": "...",
    }
    for char, replacement in replacements.items():
        raw = raw.replace(char, replacement)

    raw = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", raw)

    lines = []
    for line in raw.splitlines():
        line = line.strip()
        line = re.sub(r" {2,}", " ", line)
        lines.append(line)

    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Format detection ──────────────────────────────────────────────────────────

def _detect_format(filename: str, file_bytes: bytes) -> str:
    """
    Detect file format using magic bytes first, extension as fallback.
    Returns: "pdf" | "docx" | "unknown"
    """
    if file_bytes[:4] == b"%PDF":
        return "pdf"
    if file_bytes[:4] == b"PK\x03\x04":
        return "docx"

    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"

    return "unknown"


# ── PDF extractor ─────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        pdf_stream = io.BytesIO(file_bytes)
        parts: list = []

        with pdfplumber.open(pdf_stream) as pdf:
            total_pages = len(pdf.pages)
            if total_pages == 0:
                return "", f"[{filename}] PDF has 0 pages."

            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = ""
                try:
                    raw = page.extract_text(
                        x_tolerance=3, y_tolerance=3,
                        layout=True, x_density=7.25, y_density=13,
                    )
                    if raw:
                        page_text += raw
                except Exception as e:
                    logger.warning("[%s] Page %d text failed: %s", filename, page_num, e)

                try:
                    for table in page.extract_tables():
                        for row in table:
                            row_text = "\t".join(
                                cell.strip() if cell else "" for cell in row
                            )
                            if row_text.strip():
                                page_text += "\n" + row_text
                except Exception as e:
                    logger.debug("[%s] Page %d table failed: %s", filename, page_num, e)

                if page_text.strip():
                    parts.append(page_text)
                else:
                    logger.warning("[%s] Page %d yielded no text.", filename, page_num)

        if not parts:
            return (
                "",
                f"[{filename}] No text extracted from {total_pages} page(s). "
                "The PDF may contain only scanned images.",
            )

        cleaned = clean_text("\n\n".join(parts))
        if len(cleaned) < 50:
            return (
                "",
                f"[{filename}] Extracted text too short ({len(cleaned)} chars).",
            )

        logger.info("[%s] PDF: %d chars from %d/%d pages.", filename, len(cleaned), len(parts), total_pages)
        return cleaned, None

    except pdfplumber.exceptions.PDFSyntaxError as e:
        return "", f"[{filename}] Invalid or corrupted PDF: {e}"
    except Exception as e:
        logger.exception("[%s] Unexpected PDF extraction error", filename)
        return "", f"[{filename}] PDF extraction failed: {e}"


# ── DOCX extractor ────────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from DOCX bytes using python-docx.
    Captures: paragraphs, table cells, and text boxes (via XML fallback).
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list = []

        # 1. Body paragraphs (headings, lists, normal text)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. Tables — row by row
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    parts.append(" | ".join(row_cells))

        # 3. Text boxes (inside <w:drawing> — not exposed by python-docx API)
        drawing_texts = []
        for elem in doc.element.body.iter():
            if elem.tag == qn("a:t") and elem.text and elem.text.strip():
                drawing_texts.append(elem.text.strip())
        if drawing_texts:
            parts.append(" ".join(drawing_texts))

        if not parts:
            return (
                "",
                f"[{filename}] No text found in DOCX. "
                "The file may be empty or contain only images.",
            )

        cleaned = clean_text("\n".join(parts))
        if len(cleaned) < 50:
            return (
                "",
                f"[{filename}] Extracted text too short ({len(cleaned)} chars).",
            )

        logger.info("[%s] DOCX: %d chars, %d blocks.", filename, len(cleaned), len(parts))
        return cleaned, None

    except Exception as e:
        logger.exception("[%s] DOCX extraction error", filename)
        return "", f"[{filename}] DOCX extraction failed: {e}"


# ── Unified entry point ───────────────────────────────────────────────────────

def extract_text_from_bytes(
    file_bytes: bytes,
    filename: str = "unknown",
) -> Tuple[str, Optional[str]]:
    """
    Unified entry point — extract text from PDF or DOCX bytes.
    Auto-detects format via magic bytes. Never raises.
    Returns (text, error) — error is None on success.
    """
    if not file_bytes:
        return "", f"[{filename}] Empty file — no bytes received."

    fmt = _detect_format(filename, file_bytes)

    if fmt == "pdf":
        return _extract_pdf(file_bytes, filename)
    elif fmt == "docx":
        return _extract_docx(file_bytes, filename)
    else:
        return (
            "",
            f"[{filename}] Unsupported format. Please upload a PDF (.pdf) or Word document (.docx).",
        )


# ── Validation ────────────────────────────────────────────────────────────────

def validate_resume_file(
    filename: str,
    file_bytes: bytes,
    max_size_mb: float = MAX_SIZE_MB,
) -> Optional[str]:
    """
    Validate a resume upload (PDF or DOCX).
    Returns None if valid, or a human-readable error string.
    """
    if not filename:
        return "Filename is missing."

    lower = filename.lower()
    ext = "." + lower.rsplit(".", 1)[-1] if "." in lower else ""

    if ext not in ALLOWED_EXTENSIONS:
        return (
            f"'{filename}' is not supported. "
            "Please upload a PDF (.pdf) or Word document (.docx)."
        )

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return (
            f"'{filename}' is {size_mb:.1f} MB — exceeds the {max_size_mb:.0f} MB limit."
        )

    fmt = _detect_format(filename, file_bytes)
    if fmt == "unknown":
        return (
            f"'{filename}' does not appear to be a valid PDF or DOCX "
            "(unrecognised file header)."
        )

    return None


# Backward-compatibility alias — main.py calls validate_pdf_file()
validate_pdf_file = validate_resume_file


# ── Batch helper ──────────────────────────────────────────────────────────────

def extract_texts_from_uploads(files: list) -> list:
    """
    Process multiple resume uploads (PDF or DOCX) in one call.
    Args: list of (filename, bytes) tuples
    Returns: list of {"filename", "text", "success", "error", "char_count"}
    """
    results = []
    for filename, file_bytes in files:
        text, error = extract_text_from_bytes(file_bytes, filename)
        results.append({
            "filename":   filename,
            "text":       text,
            "success":    error is None,
            "error":      error,
            "char_count": len(text),
        })
    return results